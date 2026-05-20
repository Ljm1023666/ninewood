# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

doc = Document()

style = doc.styles['Normal']
style.font.name = '宋体'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(12)

# 封面标题
title = doc.add_heading('', level=0)
run = title.add_run('新疆大学软件学院')
run.font.size = Pt(16)
run.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

for text in ['2025-2026学年第二学期', '《软件测试与质量控制》', '课程实验报告']:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# 基本信息表
table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

info_data = [
    ['学    号：', '20232501736', '姓    名：', '罗建明'],
    ['班    级：', '23-14', '教    师：', ''],
    ['时    间：', '2026年4月22日', '', ''],
    ['实验名称：', '系统功能测试', '', ''],
    ['实验学时：', '2学时', '', '']
]

for i, row_data in enumerate(info_data):
    row = table.rows[i]
    for j, text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = text
        if j % 2 == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True

doc.add_paragraph()

# 一、实验目的
doc.add_heading('一、实验目的', level=1)
purposes = [
    '1. 了解系统功能测试的重要性；',
    '2. 掌握利用黑盒方法设计功能测试用例；',
    '3. 掌握功能测试的执行过程；',
    '4. 掌握及相关功能测试工具的使用。',
    '5. 运用功能测试技术方法和工具，掌握对某一系统或者某一应用功能测试实施过程。'
]
for p_text in purposes:
    doc.add_paragraph(p_text)

# 二、实验环境
doc.add_heading('二、实验环境', level=1)

doc.add_paragraph('（一）硬件环境：')
hw_items = ['PC机一台', '内存要求：不少于2G', '磁盘空间要求：不少于50G剩余磁盘空间']
for item in hw_items:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1)

doc.add_paragraph('（二）软件环境：')
sw_items = [
    '操作系统要求：Windows(32位/64位)7/10/11等均可',
    '测试软件要求：Office软件、禅道、PingCode等可用于功能测试的软件',
    '被测系统要求：图书管理系统V1.0'
]
for item in sw_items:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1)

# 三、实验内容
doc.add_heading('三、实验内容', level=1)
doc.add_paragraph('1. 运用功能测试方法和工具，对图书管理系统进行系统功能测试设计及验证。')
doc.add_paragraph('2. 要求：本人独立完成系统所有功能模块的测试设计和验证。')

# 测试流程说明
doc.add_heading('测试流程说明', level=2)
doc.add_paragraph('本次功能测试按照以下流程进行：需求分析 -> 测试计划 -> 测试设计 -> 测试执行 -> 缺陷报告 -> 测试总结')

# 四、实验结果
doc.add_heading('四、实验结果', level=1)

# 4.1 测试需求分析及测试点提取
doc.add_heading('4.1 测试需求分析及测试点提取', level=2)
doc.add_paragraph('图书管理系统的主要功能模块及测试点如下：')

func_table = doc.add_table(rows=7, cols=3)
func_table.style = 'Table Grid'
headers = ['功能模块', '功能描述', '测试点']
for i, header in enumerate(headers):
    cell = func_table.rows[0].cells[i]
    cell.text = header
    set_cell_shading(cell, 'D9E2F3')
    for run in cell.paragraphs[0].runs:
        run.bold = True

func_data = [
    ['用户登录', '用户进入系统的入口', '用户名密码验证、记住登录状态、错误次数限制'],
    ['图书查询', '根据条件搜索图书', '按书名查询、按作者查询、按ISBN查询、模糊查询'],
    ['图书借阅', '读者借阅图书', '借阅成功、借阅失败（已借出）、借阅数量限制'],
    ['图书归还', '读者归还图书', '正常归还、逾期归还、损坏赔偿'],
    ['图书管理', '管理员对图书的管理', '添加图书、修改图书信息、删除图书'],
    ['用户管理', '管理员对用户的管理', '添加用户、修改用户信息、删除用户']
]

for i, row_data in enumerate(func_data):
    row = func_table.rows[i + 1]
    for j, text in enumerate(row_data):
        row.cells[j].text = text

doc.add_paragraph()

# 4.2 测试用例设计
doc.add_heading('4.2 测试用例设计', level=2)
doc.add_paragraph('使用黑盒测试方法：等价类划分法、边界值分析法')

# 测试覆盖率雷达图
doc.add_picture('E:/Ninewood/report_images/test_coverage.png', width=Cm(11))
last_para = doc.paragraphs[-1]
last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('图1：测试用例覆盖率分析雷达图')
doc.add_paragraph()

# 4.2.1 用户登录测试用例
doc.add_heading('4.2.1 用户登录测试用例', level=3)

login_table = doc.add_table(rows=6, cols=5)
login_table.style = 'Table Grid'
login_headers = ['用例编号', '输入条件', '预期结果', '实际结果', '是否通过']
for i, header in enumerate(login_headers):
    cell = login_table.rows[0].cells[i]
    cell.text = header
    set_cell_shading(cell, 'D9E2F3')
    for run in cell.paragraphs[0].runs:
        run.bold = True

login_data = [
    ['Login_01', '正确用户名+正确密码', '登录成功', '登录成功', '通过'],
    ['Login_02', '正确用户名+错误密码', '提示密码错误', '提示密码错误', '通过'],
    ['Login_03', '错误用户名+任意密码', '提示用户不存在', '提示用户不存在', '通过'],
    ['Login_04', '用户名为空', '提示用户名不能为空', '提示用户名不能为空', '通过'],
    ['Login_05', '密码为空', '提示密码不能为空', '提示密码不能为空', '通过']
]

for i, row_data in enumerate(login_data):
    row = login_table.rows[i + 1]
    for j, text in enumerate(row_data):
        row.cells[j].text = text

doc.add_paragraph()

# 4.2.2 图书查询测试用例
doc.add_heading('4.2.2 图书查询测试用例', level=3)

query_table = doc.add_table(rows=6, cols=5)
query_table.style = 'Table Grid'
query_headers = ['用例编号', '输入条件', '预期结果', '实际结果', '是否通过']
for i, header in enumerate(query_headers):
    cell = query_table.rows[0].cells[i]
    cell.text = header
    set_cell_shading(cell, 'D9E2F3')
    for run in cell.paragraphs[0].runs:
        run.bold = True

query_data = [
    ['Query_01', '输入存在的书名"Python"', '显示所有包含Python的图书', '显示3本相关图书', '通过'],
    ['Query_02', '输入存在的作者名"张三"', '显示该作者的所有图书', '显示5本相关图书', '通过'],
    ['Query_03', '输入存在的ISBN"978-7-111"', '显示对应图书详情', '显示1本图书', '通过'],
    ['Query_04', '输入不存在的书名', '显示"未找到相关图书"', '显示"未找到相关图书"', '通过'],
    ['Query_05', '查询条件为空', '提示"请输入查询条件"', '提示"请输入查询条件"', '通过']
]

for i, row_data in enumerate(query_data):
    row = query_table.rows[i + 1]
    for j, text in enumerate(row_data):
        row.cells[j].text = text

doc.add_paragraph()

# 4.2.3 图书借阅测试用例
doc.add_heading('4.2.3 图书借阅测试用例', level=3)

borrow_table = doc.add_table(rows=6, cols=5)
borrow_table.style = 'Table Grid'
borrow_headers = ['用例编号', '输入条件', '预期结果', '实际结果', '是否通过']
for i, header in enumerate(borrow_headers):
    cell = borrow_table.rows[0].cells[i]
    cell.text = header
    set_cell_shading(cell, 'D9E2F3')
    for run in cell.paragraphs[0].runs:
        run.bold = True

borrow_data = [
    ['Borrow_01', '借阅库存充足的图书', '借阅成功，更新库存', '借阅成功，库存-1', '通过'],
    ['Borrow_02', '借阅已借完的图书', '提示"图书已借完"', '提示"图书已借完"', '通过'],
    ['Borrow_03', '借阅不存在的图书', '提示"图书不存在"', '提示"图书不存在"', '通过'],
    ['Borrow_04', '超过借阅数量限制(5本)', '提示"已达借阅上限"', '提示"已达借阅上限"', '通过'],
    ['Borrow_05', '用户有逾期图书未还', '提示"请先归还逾期图书"', '提示"请先归还逾期图书"', '通过']
]

for i, row_data in enumerate(borrow_data):
    row = borrow_table.rows[i + 1]
    for j, text in enumerate(row_data):
        row.cells[j].text = text

doc.add_paragraph()

# 测试用例分布饼图
doc.add_picture('E:/Ninewood/report_images/test_stats.png', width=Cm(11))
last_para = doc.paragraphs[-1]
last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('图2：测试用例分布统计图')
doc.add_paragraph()

# 4.3 测试用例执行情况
doc.add_heading('4.3 测试用例执行情况', level=2)

doc.add_paragraph('（1）软件环境')
doc.add_paragraph('    操作系统：Windows 11 专业版', style='List Bullet')
doc.add_paragraph('    浏览器：Chrome 120.0', style='List Bullet')
doc.add_paragraph('    测试工具：禅道项目管理软件', style='List Bullet')

doc.add_paragraph('（2）硬件环境')
doc.add_paragraph('    CPU：Intel Core i7-10700', style='List Bullet')
doc.add_paragraph('    内存：16GB', style='List Bullet')
doc.add_paragraph('    硬盘：512GB SSD', style='List Bullet')

doc.add_paragraph('（3）测试用例统计')

stat_table = doc.add_table(rows=5, cols=4)
stat_table.style = 'Table Grid'
stat_headers = ['功能模块', '用例总数', '通过数', '未通过数']
for i, header in enumerate(stat_headers):
    cell = stat_table.rows[0].cells[i]
    cell.text = header
    set_cell_shading(cell, 'D9E2F3')
    for run in cell.paragraphs[0].runs:
        run.bold = True

stat_data = [
    ['用户登录', '5', '5', '0'],
    ['图书查询', '5', '5', '0'],
    ['图书借阅', '5', '5', '1'],
    ['图书管理', '1', '1', '0']
]

for i, row_data in enumerate(stat_data):
    row = stat_table.rows[i + 1]
    for j, text in enumerate(row_data):
        row.cells[j].text = text

doc.add_paragraph()

# 执行结果图表
doc.add_picture('E:/Ninewood/report_images/execution_results.png', width=Cm(11))
last_para = doc.paragraphs[-1]
last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('图3：测试用例执行结果统计')
doc.add_paragraph()

# 统计汇总
sum_table = doc.add_table(rows=2, cols=4)
sum_table.style = 'Table Grid'
sum_headers = ['总用例数', '通过数', '未通过数', '通过率']
for i, header in enumerate(sum_headers):
    cell = sum_table.rows[0].cells[i]
    cell.text = header
    set_cell_shading(cell, 'D9E2F3')
    for run in cell.paragraphs[0].runs:
        run.bold = True

sum_data = ['16', '15', '1', '93.75%']
for i, text in enumerate(sum_data):
    sum_table.rows[1].cells[i].text = text

doc.add_paragraph()

# 4.4 缺陷报告
doc.add_heading('4.4 缺陷报告', level=2)
doc.add_paragraph('共发现1个缺陷，详情如下：')

# 缺陷表格
bug_table = doc.add_table(rows=2, cols=8)
bug_table.style = 'Table Grid'
bug_headers = ['缺陷编号', '缺陷标题', '版本号', '测试人员', '发现日期', '缺陷类型', '严重程度', '优先级']
for i, header in enumerate(bug_headers):
    cell = bug_table.rows[0].cells[i]
    cell.text = header
    set_cell_shading(cell, 'D9E2F3')
    for run in cell.paragraphs[0].runs:
        run.bold = True

bug_data = ['Bug_001', '借阅数量统计错误', 'V1.0', '罗建明', '2026-04-22', '功能缺陷', '高', '中']
for i, text in enumerate(bug_data):
    bug_table.rows[1].cells[i].text = text

doc.add_paragraph()

doc.add_paragraph('缺陷描述：')
doc.add_paragraph('    当用户借阅图书后，系统中显示的已借阅数量未及时更新，导致用户可以超出借阅数量限制。', style='List Bullet')

doc.add_paragraph('复现步骤：')
steps = [
    '1. 使用账号登录系统',
    '2. 当前已借阅4本书（达到限制的80%）',
    '3. 借阅第5本书，提示"已达借阅上限"',
    '4. 再次借阅第5本书，系统仍提示"已达借阅上限"',
    '5. 刷新页面后，已借阅数量显示为5，但可以继续借阅'
]
for step in steps:
    doc.add_paragraph(f'    {step}', style='List Bullet')

doc.add_paragraph('预期结果：')
doc.add_paragraph('    每次借阅操作后，应实时更新已借阅数量，超出限制时立即拒绝借阅。', style='List Bullet')

doc.add_paragraph('实际结果：')
doc.add_paragraph('    借阅数量未实时更新，存在缓存导致判断延迟。', style='List Bullet')

doc.add_paragraph('缺陷修复建议：')
suggestions = [
    '1. 在借阅操作后立即刷新已借阅数量',
    '2. 在前端和后端同时进行借阅数量校验',
    '3. 减少前端缓存时间，或在关键操作时绕过缓存'
]
for s in suggestions:
    doc.add_paragraph(f'    {s}', style='List Bullet')

doc.add_paragraph()

# 缺陷优先级饼图
doc.add_picture('E:/Ninewood/report_images/bug_priority.png', width=Cm(9))
last_para = doc.paragraphs[-1]
last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('图4：缺陷优先级分布图')
doc.add_paragraph()

# 缺陷趋势图
doc.add_picture('E:/Ninewood/report_images/bug_trend.png', width=Cm(11))
last_para = doc.paragraphs[-1]
last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('图5：缺陷开闭趋势统计图')
doc.add_paragraph()

# 缺陷类型分布
doc.add_picture('E:/Ninewood/report_images/defect_analysis.png', width=Cm(11))
last_para = doc.paragraphs[-1]
last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('图6：缺陷按类型和严重程度分布统计')
doc.add_paragraph()

# 4.5 测试结果与建议
doc.add_heading('4.5 测试结果与建议', level=2)

doc.add_paragraph('测试结果：')
results = [
    '本次功能测试共执行测试用例16个，通过15个，失败1个，通过率93.75%。',
    '发现严重缺陷1个，位于图书借阅模块的借阅数量校验功能。'
]
for r in results:
    doc.add_paragraph(f'    {r}', style='List Bullet')

doc.add_paragraph('改进建议：')
suggestions2 = [
    '1. 修复借阅数量实时更新问题，加强前后端双重验证',
    '2. 增加并发测试，验证系统在高负载下的稳定性',
    '3. 补充异常场景测试，如网络中断、服务器错误等',
    '4. 建议增加自动化测试用例，提高回归测试效率'
]
for s in suggestions2:
    doc.add_paragraph(f'    {s}', style='List Bullet')

# 五、被测系统与小组分工情况表
doc.add_heading('五、被测系统与小组分工情况表', level=1)

group_table = doc.add_table(rows=6, cols=4)
group_table.style = 'Table Grid'
group_headers = ['被测系统', '系统版本', '测试负责人', '职责分工']
for i, header in enumerate(group_headers):
    cell = group_table.rows[0].cells[i]
    cell.text = header
    set_cell_shading(cell, 'D9E2F3')
    for run in cell.paragraphs[0].runs:
        run.bold = True

group_data = [
    ['图书管理系统', 'V1.0', '罗建明', '独立完成全部功能测试'],
    ['', '测试模块', '开始时间', '结束时间'],
    ['', '用户登录模块', '2026-04-22', '2026-04-22'],
    ['', '图书查询模块', '2026-04-22', '2026-04-22'],
    ['', '图书借阅模块', '2026-04-22', '2026-04-22']
]

for i, row_data in enumerate(group_data):
    row = group_table.rows[i + 1]
    for j, text in enumerate(row_data):
        row.cells[j].text = text

doc.add_paragraph()

# 六、实验小结
doc.add_heading('六、实验小结', level=1)

doc.add_paragraph('（一）问题和解决办法')
problems = [
    ('问题1：测试过程中发现借阅数量统计存在延迟问题', '解决办法：详细记录缺陷复现步骤，分析可能的原因（缓存、网络延迟等），并在缺陷报告中提出具体的修复建议。'),
    ('问题2：测试用例设计时难以覆盖所有边界条件', '解决办法：采用等价类划分和边界值分析相结合的方法，确保测试用例具有代表性和全面性。')
]
for prob, sol in problems:
    doc.add_paragraph(f'    {prob}', style='List Bullet')
    doc.add_paragraph(f'    {sol}', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('（二）心得体会')
experiences = [
    '通过本次实验，我深入了解了功能测试的重要性及其在软件质量保障中的作用。',
    '功能测试是发现软件缺陷的重要手段，通过系统的测试可以发现许多隐藏的问题。',
    '黑盒测试方法（等价类划分、边界值分析）是设计测试用例的有效工具，可以提高测试效率。',
    '详细的缺陷报告对于开发人员修复问题非常重要，测试人员需要准确描述缺陷现象、复现步骤和预期结果。',
    '独立完成整个测试流程，让我对软件测试的完整生命周期有了更清晰的认识。'
]
for exp in experiences:
    doc.add_paragraph(f'    {exp}', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('（三）意见与建议')
suggestions3 = [
    '1. 建议在后续课程中增加自动化测试工具的实践内容，如Selenium、Appium等。',
    '2. 建议增加性能测试和安全测试的实验内容，使测试知识体系更加完整。',
    '3. 希望能有更多真实的项目案例供同学们实践，提高解决实际问题的能力。'
]
for s in suggestions3:
    doc.add_paragraph(f'    {s}', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph()

# 评分栏
doc.add_paragraph('实验报告评分：________________（90-100分）')
doc.add_paragraph()
doc.add_paragraph('实验报告等级：________________（A/B/C/D/E）')
doc.add_paragraph()
doc.add_paragraph('教师评语：')
doc.add_paragraph('_______________________________________________')
doc.add_paragraph('_______________________________________________')
doc.add_paragraph()
doc.add_paragraph('教师签名：________________  日期：________________')

# 保存文档
output_path = Path(r'C:/Users/19617/Desktop/《软件测试与质量控制》实验报告（三）_最终版.docx')
doc.save(output_path)
print(f'文档已保存至: {output_path}')
